import io
import docx
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from typing import Dict
from dependencies import get_current_user

router = APIRouter(prefix="/tools/docx", tags=["Tools"])

@router.post("/extract-text")
async def extract_docx_text(
    file: UploadFile = File(..., description="Fichier Word (.docx) à lire en mémoire"),
    current_user: dict = Depends(get_current_user)
) -> Dict[str, str]:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Le fichier doit avoir l'extension .docx")
        
    try:
        # We read the file fully into RAM
        content = await file.read()
        file_stream = io.BytesIO(content)
        
        doc = docx.Document(file_stream)
        
        extracted_text = []
        
        # 1. Extraction des paragraphes classiques (titres, textes longs)
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text.strip())
                
        # 2. Extraction des cellules de tableaux incrustés
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    extracted_text.append(" | ".join(row_data))
                    
        full_text = "\n\n".join(extracted_text)
        
        return {
            "status": "success",
            "filename": file.filename,
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
            "text": full_text
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Échec de lecture du fichier DOCX: {str(e)}")
